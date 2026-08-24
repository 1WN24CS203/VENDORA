import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_styled_document(filename):
    doc = Document()
    
    # Page setup - 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Color palette
    INK = RGBColor(26, 23, 21)        # #1A1715
    CORAL = RGBColor(232, 87, 42)     # #E8572A
    GREEN = RGBColor(45, 125, 95)     # #2D7D5F
    MUTED = RGBColor(122, 112, 103)   # #7A7067
    
    # Base Style modifications
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = INK
    
    # Header Banner
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    run_sub = p_title.add_run("COMPLETE SYSTEM MASTER MANUAL & TECHNICAL PITCH GUIDE\n")
    run_sub.font.size = Pt(10)
    run_sub.font.bold = True
    run_sub.font.color.rgb = CORAL
    
    run_main = p_title.add_run("Vendora — Vendor Registration & Management System")
    run_main.font.size = Pt(24)
    run_main.font.bold = True
    run_main.font.color.rgb = INK
    
    p_desc = doc.add_paragraph("Comprehensive documentation covering beginner analogies, business workflows, click-by-click background executions, full source code walkthroughs, database triggers, and security policies.")
    p_desc.paragraph_format.space_after = Pt(24)
    p_desc.runs[0].font.color.rgb = MUTED
    p_desc.runs[0].font.italic = True
    
    def add_h1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(22)
        h.paragraph_format.space_after = Pt(8)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = CORAL
        return h

    def add_h2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = INK
        return h

    def add_bullet(bold_prefix, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(bold_prefix + ": ")
        r1.bold = True
        r1.font.color.rgb = INK
        r2 = p.add_run(text)
        r2.font.color.rgb = INK
        return p

    def add_code_block(code_text, label=""):
        if label:
            p_lbl = doc.add_paragraph()
            p_lbl.paragraph_format.space_before = Pt(8)
            p_lbl.paragraph_format.space_after = Pt(2)
            r_l = p_lbl.add_run(f"📄 {label}")
            r_l.bold = True
            r_l.font.size = Pt(9.5)
            r_l.font.color.rgb = CORAL

        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.5)
        
        shading = parse_xml(r'<w:shd {} w:fill="1A1715"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)
        
        cp = cell.paragraphs[0]
        cp.paragraph_format.space_before = Pt(6)
        cp.paragraph_format.space_after = Pt(6)
        
        r_c = cp.add_run(code_text)
        r_c.font.name = 'Consolas'
        r_c.font.size = Pt(9)
        r_c.font.color.rgb = RGBColor(240, 235, 228)

    def add_callout(text, title="PITCH HIGHLIGHT"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.5)
        
        shading = parse_xml(r'<w:shd {} w:fill="F7F4F0"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)
        
        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(r'<w:tcBorders {}><w:left w:val="single" w:sz="24" w:space="0" w:color="E8572A"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>'.format(nsdecls('w')))
        tcPr.append(borders)
        
        cp = cell.paragraphs[0]
        cp.paragraph_format.space_before = Pt(6)
        cp.paragraph_format.space_after = Pt(6)
        
        r_t = cp.add_run(f"⚡ {title}\n")
        r_t.bold = True
        r_t.font.color.rgb = CORAL
        r_t.font.size = Pt(10)
        
        r_b = cp.add_run(text)
        r_b.font.size = Pt(10)
        r_b.font.color.rgb = INK

    # --- Section 1 ---
    add_h1("1. What is Vendora? (Explained Like You Are 5)")
    p = doc.add_paragraph("Imagine you own a big, successful toy store. To keep your shelves full of amazing toys, you buy teddy bears from Company A, remote-control cars from Company B, and board games from Company C. In business, these companies you buy from are called VENDORS (or Suppliers).")
    p.paragraph_format.space_after = Pt(8)
    
    add_bullet("Without Vendora", "You keep vendor phone numbers on sticky notes that get lost. You forget who delivered toys on time, you don't know who has verified tax details, and you risk sending money to the wrong bank account.")
    add_bullet("With Vendora", "It acts like a super-smart digital command center. It stores every vendor's phone number, email, address, tax number, bank account details, and status in one safe, clean, and organized hub.")

    add_callout("Analogy for your pitch: Vendora is like a digital vault and assistant in one. It keeps track of who is approved to do business with us, who is compliant with tax laws, and records every single management action in an unalterable audit log.", "EASY EXPLANATION")

    # --- Section 2 ---
    add_h1("2. Core Features & Business Value Matrix")
    
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Feature"
    hdr_cells[1].text = "What It Does"
    hdr_cells[2].text = "Why It Helps Businesses"
    
    for cell in hdr_cells:
        shading = parse_xml(r'<w:shd {} w:fill="1A1715"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)

    features = [
        ("3 Auth Options", "Login using Password, Email OTP, or Phone OTP", "Gives staff modern, flexible, and secure access options."),
        ("Interactive Dashboard", "Displays high-level stats, progress bars, and recent activity", "Managers can see pending approvals and compliance issues in 1 second."),
        ("5-Step Vendor Registration", "Multi-step form collecting contact info, address, bank, and docs", "Ensures zero missing data without overwhelming the user on a single long page."),
        ("Public Vendor Portal", "Public self-registration page (/vendor-register)", "Vendors fill out their own company info, saving hours of manual data entry for staff."),
        ("Vendor Categories", "Group suppliers by IT Services, Logistics, Supplies, etc.", "Allows easy sorting, filtering, and budget tracking by department."),
        ("Audit Activity Feed", "Timeline recording every approval, status change, or note", "Ensures 100% legal compliance, transparency, and accountability.")
    ]

    for feat, desc, why in features:
        row_cells = table.add_row().cells
        row_cells[0].text = feat
        row_cells[1].text = desc
        row_cells[2].text = why
        for cell in row_cells:
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            shading = parse_xml(r'<w:shd {} w:fill="FEFCFA"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- Section 3 ---
    add_h1("3. Click-by-Click Background Execution Map (Pitch Breakdown)")

    add_h2("Flow 1: User Signs Up on /register")
    add_bullet("User Action", "Enters Full Name, Email, Phone, Password, and clicks 'Create Account'.")
    add_bullet("Frontend Event", "Form submit handler in src/app/(auth)/register/page.tsx validates inputs (password >= 6 chars, match check).")
    add_bullet("Server Action", "Invokes signUpWithPassword() in src/lib/auth/actions.ts running on Next.js server.")
    add_bullet("Database Trigger", "Supabase Auth inserts into auth.users. A PostgreSQL trigger on_auth_user_created automatically executes handle_new_user() function, creating a row in public.profiles with role='admin'.")
    add_bullet("Cookie & Session", "Issues JWT access & refresh tokens set in HTTP-Only, Secure cookies via @supabase/ssr. Redirects to /dashboard.")

    add_h2("Flow 2: User Logs In on /login (3 Methods)")
    add_bullet("Password Method", "signInWithPassword() checks credentials against bcrypt hashes in auth.users, issues cookies, and redirects to dashboard.")
    add_bullet("Email / Phone OTP", "sendEmailOtp() / sendPhoneOtp() calls supabase.auth.signInWithOtp(). User enters 6-digit code (dev fallback 123456) on /verify-otp, which calls verifyOtp() to set session cookies.")

    add_h2("Flow 3: Dashboard Server Rendering (/dashboard)")
    add_bullet("Server Component", "DashboardPage() in src/app/(dashboard)/dashboard/page.tsx runs on Node.js server.")
    add_bullet("Database Queries", "Executes 2 parallel queries: (1) SELECT vendor counts by status, (2) SELECT recent activity joined with public.profiles(full_name).")
    add_bullet("Rendering", "Computes metrics in memory and streams styled HTML cards with CSS entrance animations (animate-fade-in).")

    add_h2("Flow 4: Internal Vendor Registration (5-Step Form)")
    add_bullet("Step Stepper", "Stepper.tsx updates active step state (0 to 4).")
    add_bullet("File Upload", "FileUpload.tsx validates file size <= 10MB and file extensions (.pdf, .png, .jpg).")
    add_bullet("Form Submission", "handleSubmit() calls browser client supabase.from('vendors').insert(). PostgreSQL trigger vendors_updated_at fires to set updated_at. Automatically inserts audit record into vendor_activity_log.")
    add_bullet("Cache Refresh", "router.refresh() purges Server Component cache so the vendor appears in real-time.")

    add_h2("Flow 5: Public Vendor Self-Registration Portal (/vendor-register)")
    add_bullet("Unauthenticated Access", "External vendor submits form without logging in.")
    add_bullet("Row Level Security (RLS)", "PostgreSQL evaluates RLS policy 'Public vendor self-registration'. It allows anonymous (anon) users to INSERT into vendors table ONLY IF status='pending'. Public users cannot read or update existing data.")
    add_bullet("UI Confirmation", "Renders animated SVG success screen; vendor application appears under Pending Approvals for admins.")

    add_h2("Flow 6: Vendor Status Update Modal (/vendors/[id])")
    add_bullet("Modal Mounting", "Modal.tsx locks document body scroll (overflow:hidden) and listens for Escape key.")
    add_bullet("Status Change", "Executing supabase.from('vendors').update({ status: 'approved' }).eq('id', id). Writes log entry into vendor_activity_log. Restores body scroll and refreshes UI.")

    # --- Section 4 ---
    add_h1("4. Source Code Deep-Dive & Function Details")

    add_h2("4.1 Server Actions (src/lib/auth/actions.ts)")
    p = doc.add_paragraph("Server actions handle secure user authentication directly on the Node.js server without exposing API keys to the browser:")
    p.paragraph_format.space_after = Pt(6)

    add_code_block("""export async function signUpWithPassword(formData: FormData): Promise<AuthActionResult> {
  const supabase = await createClient();
  const email = formData.get('email') as string;
  const password = formData.get('password') as string;
  const fullName = formData.get('full_name') as string;
  const phone = formData.get('phone') as string;

  const { data: signUpData, error } = await supabase.auth.signUp({
    email,
    password,
    options: { data: { full_name: fullName, phone: phone } },
  });

  if (error) return { error: error.message };

  if (signUpData.session) redirect('/dashboard');

  const { error: signInError } = await supabase.auth.signInWithPassword({ email, password });
  if (!signInError) redirect('/dashboard');

  return { success: 'Account created! You can now sign in directly.' };
}""", "signUpWithPassword() — Creates user & handles instant auto-login")

    add_code_block("""export async function sendEmailOtp(formData: FormData): Promise<AuthActionResult> {
  const supabase = await createClient();
  const email = formData.get('email') as string;

  const { error } = await supabase.auth.signInWithOtp({
    email,
    options: { shouldCreateUser: true },
  });

  if (error) {
    return { success: 'OTP sent (Dev mode: Use code 123456 for testing).', email };
  }
  return { success: 'OTP sent to your email.', email };
}""", "sendEmailOtp() — Sends OTP with dev fallback")

    add_code_block("""export async function verifyOtp(formData: FormData): Promise<AuthActionResult> {
  const supabase = await createClient();
  const token = formData.get('token') as string;
  const email = formData.get('email') as string | null;
  const phone = formData.get('phone') as string | null;

  if (email) {
    const { error } = await supabase.auth.verifyOtp({ email, token, type: 'email' });
    if (error && token === '123456') redirect('/dashboard');
    if (error) return { error: error.message };
  }
  redirect('/dashboard');
}""", "verifyOtp() — Verifies 6-digit code or dev bypass")

    add_h2("4.2 Session Middleware (src/lib/supabase/middleware.ts)")
    p = doc.add_paragraph("Middleware intercepts every web request, refreshes expired JWT tokens, and redirects unauthenticated users away from private dashboard pages:")
    p.paragraph_format.space_after = Pt(6)

    add_code_block("""export async function updateSession(request: NextRequest) {
  let supabaseResponse = NextResponse.next({ request });
  const supabase = createServerClient(
    process.env.NEXT_PUBLIC_SUPABASE_URL!,
    process.env.NEXT_PUBLIC_SUPABASE_ANON_KEY!,
    {
      cookies: {
        getAll() { return request.cookies.getAll(); },
        setAll(cookiesToSet) {
          cookiesToSet.forEach(({ name, value, options }) =>
            supabaseResponse.cookies.set(name, value, options)
          );
        },
      },
    }
  );

  const { data: { user } } = await supabase.auth.getUser();

  const isAuthRoute = request.nextUrl.pathname.startsWith('/login') ||
    request.nextUrl.pathname.startsWith('/register') ||
    request.nextUrl.pathname.startsWith('/verify-otp');
  const isDashboardRoute = request.nextUrl.pathname.startsWith('/dashboard') ||
    request.nextUrl.pathname.startsWith('/vendors') ||
    request.nextUrl.pathname.startsWith('/categories') ||
    request.nextUrl.pathname.startsWith('/settings');

  if (!user && isDashboardRoute) {
    const url = request.nextUrl.clone();
    url.pathname = '/login';
    return NextResponse.redirect(url);
  }

  if (user && isAuthRoute) {
    const url = request.nextUrl.clone();
    url.pathname = '/dashboard';
    return NextResponse.redirect(url);
  }

  return supabaseResponse;
}""", "updateSession() — Route Protection & Session Refresh")

    add_h2("4.3 Database Triggers & RLS Policies (supabase/schema.sql)")
    p = doc.add_paragraph("PostgreSQL functions, triggers, and Row Level Security policies enforcing data integrity and security directly inside the database:")
    p.paragraph_format.space_after = Pt(6)

    add_code_block("""-- Auto-create profile on signup
CREATE OR REPLACE FUNCTION public.handle_new_user()
RETURNS TRIGGER AS $$
BEGIN
  INSERT INTO public.profiles (id, full_name, role)
  VALUES (
    NEW.id,
    COALESCE(NEW.raw_user_meta_data ->> 'full_name', NEW.email),
    'admin'
  );
  RETURN NEW;
END;
$$ LANGUAGE plpgsql SECURITY DEFINER;

CREATE OR REPLACE TRIGGER on_auth_user_created
  AFTER INSERT ON auth.users
  FOR EACH ROW EXECUTE FUNCTION public.handle_new_user();""", "PostgreSQL Trigger — Auto-Profile Creation")

    add_code_block("""-- Public vendor self-registration policy
CREATE POLICY "Public vendor self-registration"
  ON public.vendors FOR INSERT
  TO anon
  WITH CHECK (status = 'pending');""", "PostgreSQL RLS Policy — Restricts Anonymous Self-Registration to Pending Status")

    # --- Section 5 ---
    add_h1("5. Code Directory Blueprint")
    add_bullet("src/app/(auth)", "Contains login, register, and verify-otp pages with split-screen editorial design.")
    add_bullet("src/app/(dashboard)", "Contains Dashboard, All Vendors table, Add Vendor form, Vendor Detail page, Categories, and Settings.")
    add_bullet("src/app/vendor-register", "Standalone public registration page for self-onboarding.")
    add_bullet("src/lib/supabase", "Handles browser and server database connections plus session middleware.")
    add_bullet("src/lib/auth/actions.ts", "Server Actions for user signup, password signin, OTP generation, and signout.")
    add_bullet("src/components/ui", "13 bespoke UI components (Buttons, Inputs, Modals, Badges, Stepper, Toast, ActivityFeed, SearchBar, FileUpload).")
    add_bullet("supabase/schema.sql", "Complete SQL database setup script with 5 tables, RLS policies, triggers, and default categories.")

    # Save document
    doc.save(filename)
    print(f"Document successfully created: {filename}")

if __name__ == '__main__':
    create_styled_document("c:\\Users\\prath\\OneDrive - bmsce.ac.in\\Desktop\\Namith Project\\Vendora_Complete_Master_Guide.docx")
